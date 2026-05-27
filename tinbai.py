import streamlit as st
import pandas as pd
from supabase import create_client
from datetime import datetime
import io
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import plotly.express as px
import plotly.graph_objects as go
import unicodedata
import re

# ─────────────────────────────────────────────
# 1. CẤU HÌNH TRANG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Quản lý Tin bài – Ban Tuyên giáo",
    page_icon="📰",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
# 2. KẾT NỐI SUPABASE
# ─────────────────────────────────────────────
try:
    URL        = st.secrets["SUPABASE_URL"]
    KEY        = st.secrets["SUPABASE_KEY"]
    ADMIN_PASS = st.secrets.get("ADMIN_PASS", "141983")
    supabase   = create_client(URL, KEY)
except Exception as e:
    st.error("⚠️ Lỗi kết nối cơ sở dữ liệu. Vui lòng kiểm tra cấu hình Secrets!")
    st.stop()

try:
    supabase.table("thong_ke_truy_cap").insert({"ten_app": "Đăng ký Tin bài"}).execute()
except Exception:
    pass

# ─────────────────────────────────────────────
# 3. CSS – PHONG CÁCH HÀNH CHÍNH HIỆN ĐẠI
# ─────────────────────────────────────────────
CUSTOM_CSS = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Be+Vietnam+Pro:wght@300;400;500;600;700;800&display=swap');

    /* ── Nền & font chữ ── */
    html, body, [data-testid="stAppViewContainer"] {
        background-color: #EEF2F7;
        font-family: 'Be Vietnam Pro', 'Segoe UI', sans-serif;
    }

    /* ── Header banner ── */
    .header-banner {
        background: linear-gradient(120deg, #0A2744 0%, #0D3A6E 55%, #1155A0 100%);
        border-bottom: 5px solid #B22222;
        padding: 26px 36px 20px;
        border-radius: 12px;
        color: white;
        text-align: center;
        box-shadow: 0 6px 24px rgba(10,39,68,0.22);
        margin-bottom: 28px;
        position: relative;
        overflow: hidden;
    }
    .header-banner::before {
        content: "";
        position: absolute; top: 0; left: 0; right: 0; bottom: 0;
        background: url("data:image/svg+xml,%3Csvg width='60' height='60' viewBox='0 0 60 60' xmlns='http://www.w3.org/2000/svg'%3E%3Cg fill='none' fill-rule='evenodd'%3E%3Cg fill='%23ffffff' fill-opacity='0.03'%3E%3Cpath d='M36 34v-4h-2v4h-4v2h4v4h2v-4h4v-2h-4zm0-30V0h-2v4h-4v2h4v4h2V6h4V4h-4zM6 34v-4H4v4H0v2h4v4h2v-4h4v-2H6zM6 4V0H4v4H0v2h4v4h2V6h4V4H6z'/%3E%3C/g%3E%3C/g%3E%3C/svg%3E");
        pointer-events: none;
    }
    .header-banner .org-name {
        font-size: 0.82rem;
        font-weight: 500;
        color: rgba(255,255,255,0.7);
        letter-spacing: 0.12em;
        text-transform: uppercase;
        margin-bottom: 8px;
    }
    .header-banner h1 {
        font-size: 1.55rem;
        font-weight: 800;
        letter-spacing: 0.04em;
        text-transform: uppercase;
        margin: 0 0 8px 0;
        text-shadow: 0 2px 8px rgba(0,0,0,0.3);
        line-height: 1.3;
    }
    .header-banner .subtitle {
        font-size: 0.95rem;
        font-weight: 400;
        color: #FFD700;
        margin: 0;
        letter-spacing: 0.02em;
    }
    .header-banner .star-row {
        font-size: 1.1rem;
        color: #FFD700;
        letter-spacing: 8px;
        display: block;
        margin-bottom: 12px;
        opacity: 0.9;
    }
    .header-banner .red-line {
        width: 60px; height: 3px;
        background: #FFD700;
        margin: 12px auto 0;
        border-radius: 2px;
    }

    /* ── Sidebar ── */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #071A2E 0%, #0A2744 100%) !important;
        border-right: 3px solid #B22222;
    }
    [data-testid="stSidebar"] * { color: #C8D8EA !important; }
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {
        color: #FFD700 !important;
        border-bottom: 1px solid rgba(255,215,0,0.25);
        padding-bottom: 8px;
        font-weight: 700 !important;
    }

    /* Ô nhập mật khẩu – nền trắng, chữ đen */
    [data-testid="stSidebar"] .stTextInput input {
        background: #FFFFFF !important;
        border: 1.5px solid #3A7BD5 !important;
        color: #1a1a1a !important;
        border-radius: 6px !important;
        font-weight: 500 !important;
        caret-color: #1a1a1a !important;
    }
    [data-testid="stSidebar"] .stTextInput label {
        color: #C8D8EA !important;
        font-weight: 600 !important;
    }

    /* ── Tab header ── */
    [data-testid="stTabs"] [data-baseweb="tab-list"] {
        background: #FFFFFF;
        border-radius: 10px 10px 0 0;
        padding: 5px 10px 0;
        border-bottom: 2px solid #0A2744;
        gap: 3px;
        box-shadow: 0 2px 8px rgba(10,39,68,0.08);
    }
    [data-testid="stTabs"] [data-baseweb="tab"] {
        font-weight: 600;
        color: #555;
        border-radius: 8px 8px 0 0;
        padding: 10px 22px;
        font-size: 0.88rem;
        letter-spacing: 0.01em;
    }
    [data-testid="stTabs"] [aria-selected="true"] {
        background: #0A2744 !important;
        color: white !important;
        border-bottom: 3px solid #FFD700 !important;
    }

    /* ── Nút bấm chính ── */
    .stButton > button {
        background: linear-gradient(135deg, #0A2744 0%, #1155A0 100%);
        color: white;
        border: none;
        border-radius: 7px;
        font-weight: 700;
        padding: 0.52rem 1.5rem;
        font-size: 0.88rem;
        transition: all 0.22s ease;
        letter-spacing: 0.02em;
        box-shadow: 0 3px 10px rgba(10,39,68,0.2);
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #B22222 0%, #DC143C 100%);
        box-shadow: 0 5px 16px rgba(178,34,34,0.3);
        transform: translateY(-1px);
    }

    /* ── Nút tải xuống Word ── */
    [data-testid="stDownloadButton"] > button {
        background: linear-gradient(135deg, #0F5132 0%, #198754 100%) !important;
        color: white !important;
        border-radius: 7px !important;
        font-weight: 700 !important;
        border: none !important;
        box-shadow: 0 3px 10px rgba(15,81,50,0.25) !important;
    }
    [data-testid="stDownloadButton"] > button:hover {
        background: linear-gradient(135deg, #0A3D22 0%, #146C43 100%) !important;
        transform: translateY(-1px) !important;
    }

    /* ── Form container ── */
    [data-testid="stForm"] {
        background: white;
        border-radius: 12px;
        padding: 28px;
        border: 1px solid #D0DFF0;
        box-shadow: 0 3px 16px rgba(10,39,68,0.08);
    }

    /* ── Metric cards ── */
    [data-testid="stMetric"] {
        background: white;
        border-radius: 12px;
        padding: 18px 22px;
        border-left: 5px solid #0A2744;
        box-shadow: 0 3px 12px rgba(10,39,68,0.09);
    }
    [data-testid="stMetricValue"] {
        color: #B22222 !important;
        font-size: 2rem !important;
        font-weight: 800 !important;
    }
    [data-testid="stMetricLabel"] {
        color: #0A2744 !important;
        font-weight: 600 !important;
        font-size: 0.82rem !important;
    }

    /* ── Section heading ── */
    .section-title {
        color: #0A2744;
        font-weight: 800;
        border-left: 5px solid #B22222;
        padding: 6px 0 6px 14px;
        margin: 0 0 16px 0;
        font-size: 1.05rem;
        letter-spacing: 0.02em;
    }

    /* ── Section divider ── */
    .section-divider {
        height: 3px;
        background: linear-gradient(90deg, #0A2744 0%, #B22222 50%, #FFD700 100%);
        border-radius: 2px;
        margin: 20px 0 24px;
        opacity: 0.7;
    }

    /* ── Info box ── */
    .info-box {
        background: #EBF3FF;
        border: 1px solid #BDD4F0;
        border-left: 4px solid #1155A0;
        border-radius: 8px;
        padding: 12px 18px;
        font-size: 0.88rem;
        color: #1a3a6e;
        margin: 10px 0;
    }

    /* ── Radio button ── */
    [data-testid="stRadio"] > label { font-weight: 600; color: #0A2744; }
    [data-testid="stRadio"] [data-testid="stMarkdown"] p { font-weight: 500; }

    /* ── Expander ── */
    [data-testid="stExpander"] {
        border: 1px solid #D0DFF0 !important;
        border-radius: 10px !important;
        background: white !important;
    }
    [data-testid="stExpander"] summary {
        font-weight: 700;
        color: #0A2744;
    }

    /* ── Dataframe / Table ── */
    [data-testid="stDataFrame"] { border-radius: 10px; overflow: hidden; }

    /* ── Select box / text input chung ── */
    .stSelectbox select, .stTextInput input, .stTextArea textarea {
        border-radius: 7px !important;
        border: 1.5px solid #C8D8EA !important;
        font-family: 'Be Vietnam Pro', sans-serif !important;
    }
    .stSelectbox select:focus, .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: #1155A0 !important;
        box-shadow: 0 0 0 3px rgba(17,85,160,0.12) !important;
    }

    /* ── Alert ── */
    [data-testid="stAlert"] { border-radius: 9px; }

    /* ── Admin badge ── */
    .badge-admin {
        display: inline-block;
        background: #FFD700;
        color: #0A2744;
        font-weight: 800;
        padding: 3px 14px;
        border-radius: 20px;
        font-size: 0.74rem;
        letter-spacing: 0.06em;
        text-transform: uppercase;
    }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# ─────────────────────────────────────────────
# 4. CÁC HÀM TIỆN ÍCH
# ─────────────────────────────────────────────
def lam_sach_ten_file(ten_file: str) -> str:
    ten_file = unicodedata.normalize("NFKD", ten_file).encode("ASCII", "ignore").decode("utf-8")
    ten_file = ten_file.replace(" ", "_")
    ten_file = re.sub(r"[^\w.\-]", "", ten_file)
    return ten_file


def them_hyperlink_vao_cell(cell, text: str, url: str):
    """Chèn hyperlink vào ô bảng Word."""
    if not url or url.strip() in ("", "nan"):
        cell.text = text
        return
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text if text and text.lower() != "nan" else url
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def tao_file_word(df: pd.DataFrame, ngay_thang: str) -> bytes:
    """Tạo file Word tổng hợp, tự động gắn link cho cả bài mới lẫn bài đăng lại."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(1.5))

    # ── Tiêu đề ──
    heading = doc.add_heading("BẢNG TỔNG HỢP ĐĂNG KÝ TIN BÀI", 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(10, 39, 68)
        run.font.size = Pt(16)
        run.font.bold = True

    p_sub = doc.add_paragraph("Ban Tuyên giáo và Dân vận Tỉnh ủy Tuyên Quang")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.runs[0]
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(10, 39, 68)
    r_sub.font.size = Pt(11)

    p_ngay = doc.add_paragraph(f"Ngày tổng hợp: {ngay_thang}   |   Tổng số bài: {len(df)}")
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ngay = p_ngay.runs[0]
    r_ngay.font.italic = True
    r_ngay.font.color.rgb = RGBColor(100, 100, 100)
    r_ngay.font.size = Pt(10)

    doc.add_paragraph()

    # ── Bảng ──
    headers    = ["STT", "Tiêu đề bài viết", "Người gửi", "Loại bài", "Nguồn / Link đính kèm", "Đề xuất MXH", "Người đăng"]
    col_widths = [Cm(1.0), Cm(8.0), Cm(3.2), Cm(2.8), Cm(5.5), Cm(3.5), Cm(3.5)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for i, col in enumerate(table.columns):
        col.width = col_widths[i]

    hdr_row = table.rows[0]
    hdr_row.height = Cm(1.0)
    for i, (cell, text) in enumerate(zip(hdr_row.cells, headers)):
        cell.width = col_widths[i]
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0A2744")
        tc_pr.append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, row in enumerate(df.itertuples(), 1):
        rc = table.add_row().cells

        # STT
        rc[0].text = str(idx)
        rc[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Tiêu đề
        rc[1].text = str(row.tieu_de) if pd.notna(row.tieu_de) else ""

        # Người gửi
        rc[2].text = str(row.nguoi_gui) if pd.notna(row.nguoi_gui) else ""

        # Loại bài
        nguon_str = str(row.nguon_tin) if pd.notna(row.nguon_tin) else ""
        rc[3].text = nguon_str
        rc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Nguồn / Link – tự động gắn link cả bài mới lẫn bài đăng lại
        duong_dan = str(getattr(row, "duong_dan", "") or "")
        duong_dan = duong_dan.strip() if duong_dan.lower() not in ("nan", "") else ""

        is_suu_tam = any(kw in nguon_str.lower() for kw in ["sưu tầm", "đăng lại"])

        if duong_dan:
            # Lấy link đầu tiên nếu có nhiều
            first_link = duong_dan.splitlines()[0].strip()
            # Bài mới: hiển thị tên file (phần cuối URL); bài sưu tầm: hiển thị tên nguồn hoặc link
            if is_suu_tam:
                hien_thi = nguon_str if nguon_str and nguon_str.lower() != "nan" else first_link
            else:
                # Bài mới: lấy tên file từ URL
                hien_thi = first_link.split("/")[-1].split("?")[0] or first_link
                hien_thi = hien_thi[:60] + "…" if len(hien_thi) > 60 else hien_thi
            them_hyperlink_vao_cell(rc[4], hien_thi, first_link)
        else:
            rc[4].text = nguon_str if is_suu_tam else "—"

        # Đề xuất MXH
        mxh = ["Web"]
        if getattr(row, "dang_facebook", False): mxh.append("Facebook")
        if getattr(row, "dang_zalo",    False): mxh.append("Zalo OA")
        rc[5].text = " | ".join(mxh)
        rc[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Người đăng
        nguoi_dang = getattr(row, "nguoi_dang", "")
        rc[6].text = str(nguoi_dang) if pd.notna(nguoi_dang) else ""

        # Độ rộng ô & cỡ chữ
        for i, cell in enumerate(rc):
            cell.width = col_widths[i]
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

        # Màu xen kẽ
        if idx % 2 == 0:
            for cell in rc:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EBF3FF")
                tc_pr.append(shd)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _xu_ly_gui(nguoi_gui, la_viet_moi, nguon_label, ghi_chu,
               tieu_de_viet_moi, file_uploads, tieu_de_suu_tam, link_suu_tam):
    """Xử lý logic gửi đăng ký tin bài."""
    if not nguoi_gui.strip():
        st.error("⛔ Vui lòng điền Họ và tên trước khi gửi!")
        return

    nguoi_gui_clean = nguoi_gui.strip()
    ghi_chu_clean   = ghi_chu.strip() if ghi_chu else ""

    if la_viet_moi:
        if not tieu_de_viet_moi.strip():
            st.error("⛔ Vui lòng nhập Tiêu đề bài viết!")
            return
        if not file_uploads:
            st.error("⛔ Chọn 'Viết mới' nhưng chưa tải file nào lên!")
            return

        links = []
        for f in file_uploads:
            ts        = datetime.now().strftime("%Y%m%d%H%M%S")
            raw_name  = f"{ts}_{f.name}"
            safe_name = lam_sach_ten_file(raw_name)
            try:
                supabase.storage.from_("tin_bai").upload(safe_name, f.read())
                url = supabase.storage.from_("tin_bai").get_public_url(safe_name)
                links.append(url)
            except Exception as e:
                st.error(f"Lỗi tải file {f.name}: {e}")

        if links:
            supabase.table("dang_ky_tin_bai").insert({
                "nguoi_gui": nguoi_gui_clean,
                "tieu_de":   tieu_de_viet_moi.strip(),
                "nguon_tin": nguon_label,
                "duong_dan": "\n".join(links),
                "ghi_chu":   ghi_chu_clean,
            }).execute()
            st.success(f"✅ Đã gửi thành công bài viết mới (kèm {len(links)} file đính kèm)!")
        else:
            st.error("⛔ Không upload được file nào. Vui lòng thử lại!")
    else:
        if not tieu_de_suu_tam.strip():
            st.error("⛔ Vui lòng nhập tiêu đề bài sưu tầm!")
            return
        if not link_suu_tam.strip():
            st.error("⛔ Vui lòng nhập đường dẫn bài sưu tầm!")
            return
        try:
            supabase.table("dang_ky_tin_bai").insert({
                "nguoi_gui": nguoi_gui_clean,
                "tieu_de":   tieu_de_suu_tam.strip(),
                "nguon_tin": nguon_label,
                "duong_dan": link_suu_tam.strip(),
                "ghi_chu":   ghi_chu_clean,
            }).execute()
            st.success("✅ Đã gửi đăng ký bài sưu tầm/đăng lại thành công!")
        except Exception as e:
            st.error(f"Lỗi khi gửi: {e}")


# ─────────────────────────────────────────────
# 5. HEADER
# ─────────────────────────────────────────────
st.markdown("""
    <div class="header-banner">
        <div class="org-name">Tỉnh ủy Tuyên Quang</div>
        <span class="star-row">★ ★ ★ ★ ★</span>
        <h1>Hệ thống Quản lý Tin bài<br>Đăng Trang Thông tin Điện tử</h1>
        <div class="subtitle">Ban Tuyên giáo và Dân vận Tỉnh ủy Tuyên Quang</div>
        <div class="red-line"></div>
    </div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# 6. SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 🔐 Quản trị hệ thống")
    mat_khau = st.text_input("Mật khẩu quản trị:", type="password", placeholder="Nhập mật khẩu…")
    is_admin = mat_khau == ADMIN_PASS
    if is_admin:
        st.success("✔ Đã xác thực quản trị viên")
    elif mat_khau:
        st.error("✘ Mật khẩu không đúng")

    st.markdown("---")
    st.markdown("""
        <div style='font-size:0.78rem; color:#7B9EC4; line-height:1.9;'>
            <b style='color:#FFD700; font-size:0.82rem;'>📌 Hướng dẫn nhanh:</b><br>
            ✏️ Tab 1: Cán bộ đăng ký tin bài<br>
            📋 Tab 2: Quản trị duyệt & xuất Word<br>
            📊 Tab 3: Thống kê biểu đồ<br>
            <br>
            <b style='color:#FFD700; font-size:0.82rem;'>🔧 Quyền quản trị:</b><br>
            • Chỉnh sửa thông tin bài đã gửi<br>
            • Xóa bài đăng ký nhầm<br>
            • Lưu trạng thái MXH & Người đăng<br>
            • Xuất file Word trình duyệt<br>
        </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────
# 7. TABS CHÍNH
# ─────────────────────────────────────────────
tab1, tab2, tab3 = st.tabs([
    "✍️  Đăng ký tin bài",
    "📋  Tổng hợp & Duyệt",
    "📊  Thống kê & Biểu đồ",
])

# ══════════════════════════════════════════════
# TAB 1 – ĐĂNG KÝ TIN BÀI
# ══════════════════════════════════════════════
with tab1:
    st.markdown('<p class="section-title">📝 Gửi đăng ký bài viết</p>', unsafe_allow_html=True)

    st.markdown("""
        <div class="info-box">
            💡 <b>Hướng dẫn:</b> Chọn loại bài (Viết mới hoặc Sưu tầm/Đăng lại), 
            điền đầy đủ thông tin và nhấn <b>Gửi đăng ký</b>. 
            Các trường có dấu <b>*</b> là bắt buộc.
        </div>
    """, unsafe_allow_html=True)

    nguon = st.radio(
        "📂 **Loại tin bài:**",
        ["✏️  Bài tự viết (mới)", "🔗  Đề nghị đăng lại / Sưu tầm"],
        horizontal=True,
    )
    la_viet_moi = "tự viết" in nguon

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    with st.form("form_dang_ky", clear_on_submit=True):
        nguoi_gui = st.text_input(
            "👤 Họ và tên người gửi *",
            placeholder="Nhập họ và tên đầy đủ…"
        )

        tieu_de_viet_moi = ""
        file_uploads     = []
        tieu_de_suu_tam  = ""
        link_suu_tam     = ""

        if la_viet_moi:
            st.markdown("---")
            st.markdown("**📄 Thông tin bài viết mới**")
            tieu_de_viet_moi = st.text_input(
                "Tiêu đề bài viết *",
                placeholder="Ví dụ: Infographic về chuyển đổi số năm 2025…"
            )
            file_uploads = st.file_uploader(
                "📎 Tải lên file đính kèm (có thể chọn nhiều):",
                type=["doc", "docx", "pdf", "png", "jpg", "jpeg"],
                accept_multiple_files=True,
                help="Hỗ trợ: .doc, .docx, .pdf, .png, .jpg"
            )
        else:
            st.markdown("---")
            st.markdown("**🔗 Thông tin bài sưu tầm / đăng lại**")
            tieu_de_suu_tam = st.text_input(
                "Tiêu đề bài sưu tầm *",
                placeholder="Nhập tiêu đề bài viết gốc…"
            )
            link_suu_tam = st.text_input(
                "Đường dẫn nguồn (URL) *",
                placeholder="https://baotuyenquang.com.vn/..."
            )

        ghi_chu = st.text_area(
            "💬 Ghi chú thêm:",
            placeholder="Thông tin bổ sung, yêu cầu đặc biệt… (không bắt buộc)",
            height=80
        )

        btn_gui = st.form_submit_button(
            "📨  Gửi đăng ký tin bài",
            use_container_width=True,
            type="primary"
        )
        if btn_gui:
            _xu_ly_gui(
                nguoi_gui        = nguoi_gui,
                la_viet_moi      = la_viet_moi,
                nguon_label      = "Viết mới" if la_viet_moi else "Đề nghị đăng lại (Sưu tầm)",
                ghi_chu          = ghi_chu,
                tieu_de_viet_moi = tieu_de_viet_moi,
                file_uploads     = file_uploads,
                tieu_de_suu_tam  = tieu_de_suu_tam,
                link_suu_tam     = link_suu_tam,
            )

# ══════════════════════════════════════════════
# TAB 2 – TỔNG HỢP & DUYỆT
# ══════════════════════════════════════════════
with tab2:
    if not is_admin:
        st.warning("🔒 Vui lòng nhập mật khẩu Quản trị ở thanh bên trái để sử dụng chức năng này.")
        st.stop()

    st.markdown('<p class="section-title">📋 Bảng tổng hợp & Phê duyệt</p>', unsafe_allow_html=True)
    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    c_date, c_info = st.columns([2, 3])
    with c_date:
        ngay_xem = st.date_input("📅 Chọn ngày:", datetime.now().date())
    ngay_xem_str = ngay_xem.strftime("%d/%m/%Y")

    res = supabase.table("dang_ky_tin_bai").select("*").eq("ngay_dang_ky", ngay_xem.isoformat()).execute()

    if not res.data:
        st.info(f"ℹ️ Không có tin bài nào được đăng ký ngày **{ngay_xem_str}**.")
    else:
        df_ngay = pd.DataFrame(res.data)
        for col in ("nguoi_dang", "dang_facebook", "dang_zalo"):
            if col not in df_ngay.columns:
                df_ngay[col] = False if col.startswith("dang") else ""
        df_ngay["nguoi_dang"]    = df_ngay["nguoi_dang"].fillna("")
        df_ngay["dang_facebook"] = df_ngay["dang_facebook"].fillna(False)
        df_ngay["dang_zalo"]     = df_ngay["dang_zalo"].fillna(False)

        m1, m2, m3 = st.columns(3)
        m1.metric("📰 Tổng bài", len(df_ngay))
        m2.metric("✏️ Tự viết",  len(df_ngay[df_ngay["nguon_tin"].str.contains("Viết mới|tự viết", case=False, na=False)]))
        m3.metric("🔗 Sưu tầm", len(df_ngay[df_ngay["nguon_tin"].str.contains("Sưu tầm|đăng lại", case=False, na=False)]))

        st.markdown(f"<br>", unsafe_allow_html=True)

        # ── Data editor ──
        edited_df = st.data_editor(
            df_ngay,
            column_config={
                "id":            st.column_config.NumberColumn("ID", disabled=True),
                "tieu_de":       st.column_config.TextColumn("📝 Tiêu đề"),
                "nguoi_gui":     st.column_config.TextColumn("👤 Người gửi"),
                "nguon_tin":     st.column_config.TextColumn("Loại bài", disabled=True),
                "dang_facebook": st.column_config.CheckboxColumn("📘 Facebook", default=False),
                "dang_zalo":     st.column_config.CheckboxColumn("💬 Zalo", default=False),
                "duong_dan":     st.column_config.TextColumn("🔗 Link / File"),
                "nguoi_dang":    st.column_config.TextColumn("👤 Người đăng"),
                "ghi_chu":       st.column_config.TextColumn("💬 Ghi chú"),
                "ngay_dang_ky":  st.column_config.DateColumn("📅 Ngày", disabled=True),
            },
            hide_index=True,
            use_container_width=True,
            num_rows="dynamic",
        )

        st.markdown("")
        col_luu, col_word, col_space = st.columns([2, 2, 1])

        # ── Nút Lưu ──
        with col_luu:
            if st.button("💾  Lưu thay đổi (MXH & Người đăng)", use_container_width=True):
                ok = err = 0
                for _, row in edited_df.iterrows():
                    row_id = row.get("id")
                    if row_id:
                        try:
                            supabase.table("dang_ky_tin_bai").update({
                                "tieu_de":       str(row.get("tieu_de", "")),
                                "nguoi_gui":     str(row.get("nguoi_gui", "")),
                                "dang_facebook": bool(row.get("dang_facebook", False)),
                                "dang_zalo":     bool(row.get("dang_zalo", False)),
                                "nguoi_dang":    str(row.get("nguoi_dang", "")),
                                "ghi_chu":       str(row.get("ghi_chu", "")),
                            }).eq("id", row_id).execute()
                            ok += 1
                        except Exception:
                            err += 1
                if ok:  st.success(f"✅ Đã lưu {ok} tin bài thành công!")
                if err: st.error(f"⚠️ {err} bài gặp lỗi khi lưu.")
                st.rerun()

        # ── Nút Xuất Word ──
        with col_word:
            word_bytes = tao_file_word(edited_df, ngay_xem_str)
            st.download_button(
                "📥  Xuất file Word trình duyệt",
                data=word_bytes,
                file_name=f"TinBai_{ngay_xem.strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    # ══ Xóa bài chọn lọc ══
    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
    with st.expander("🗑️  Xóa bài đăng ký nhầm (chọn lọc)", expanded=False):
        st.markdown("""
            <div class="info-box">
                ⚠️ Chọn đúng bài cần xóa trước khi nhấn nút. Thao tác không thể hoàn tác.
            </div>
        """, unsafe_allow_html=True)

        ngay_xoa = st.date_input("📅 Chọn ngày cần xóa bài:", datetime.now().date(), key="date_xoa")
        res_xoa  = supabase.table("dang_ky_tin_bai").select("id, tieu_de, nguoi_gui, nguon_tin").eq(
            "ngay_dang_ky", ngay_xoa.isoformat()
        ).execute()

        if not res_xoa.data:
            st.info("Không có bài nào trong ngày này.")
        else:
            df_xoa = pd.DataFrame(res_xoa.data)
            # Tạo nhãn chọn
            df_xoa["label"] = df_xoa.apply(
                lambda r: f"[{r['id']}] {r['tieu_de'][:50]}… — {r['nguoi_gui']} ({r['nguon_tin']})",
                axis=1
            )
            chon_xoa = st.multiselect(
                "Chọn bài cần xóa:",
                options=df_xoa["id"].tolist(),
                format_func=lambda x: df_xoa.loc[df_xoa["id"] == x, "label"].values[0],
            )

            if chon_xoa:
                st.warning(f"Sắp xóa **{len(chon_xoa)} bài**. Xác nhận?")
                if st.button("⚠️ Xác nhận xóa các bài đã chọn", key="btn_xoa_chon"):
                    xoa_ok = xoa_err = 0
                    for bid in chon_xoa:
                        try:
                            supabase.table("dang_ky_tin_bai").delete().eq("id", bid).execute()
                            xoa_ok += 1
                        except Exception:
                            xoa_err += 1
                    if xoa_ok:  st.success(f"✅ Đã xóa {xoa_ok} bài!")
                    if xoa_err: st.error(f"⚠️ {xoa_err} bài lỗi khi xóa.")
                    st.rerun()

    # ══ Sửa từng bài ══
    with st.expander("✏️  Sửa thông tin bài đăng ký", expanded=False):
        st.markdown("""
            <div class="info-box">
                📝 Chọn bài cần sửa từ danh sách, chỉnh tiêu đề / họ tên người gửi rồi nhấn <b>Lưu sửa đổi</b>.
            </div>
        """, unsafe_allow_html=True)

        ngay_sua = st.date_input("📅 Chọn ngày:", datetime.now().date(), key="date_sua")
        res_sua  = supabase.table("dang_ky_tin_bai").select("*").eq(
            "ngay_dang_ky", ngay_sua.isoformat()
        ).execute()

        if not res_sua.data:
            st.info("Không có bài nào trong ngày này.")
        else:
            df_sua = pd.DataFrame(res_sua.data)
            df_sua["label"] = df_sua.apply(
                lambda r: f"[{r['id']}] {str(r['tieu_de'])[:50]} — {r['nguoi_gui']}",
                axis=1
            )
            id_chon = st.selectbox(
                "Chọn bài cần sửa:",
                options=df_sua["id"].tolist(),
                format_func=lambda x: df_sua.loc[df_sua["id"] == x, "label"].values[0],
                key="select_sua"
            )

            bai_chon = df_sua[df_sua["id"] == id_chon].iloc[0]

            with st.form("form_sua_bai"):
                col_a, col_b = st.columns(2)
                with col_a:
                    ten_moi = st.text_input(
                        "✏️ Tiêu đề bài viết:",
                        value=str(bai_chon["tieu_de"]) if pd.notna(bai_chon["tieu_de"]) else ""
                    )
                with col_b:
                    nguoi_gui_moi = st.text_input(
                        "👤 Họ và tên người gửi:",
                        value=str(bai_chon["nguoi_gui"]) if pd.notna(bai_chon["nguoi_gui"]) else ""
                    )

                link_moi = st.text_input(
                    "🔗 Đường dẫn / Link:",
                    value=str(bai_chon.get("duong_dan", "")) if pd.notna(bai_chon.get("duong_dan", "")) else ""
                )
                ghi_chu_moi = st.text_area(
                    "💬 Ghi chú:",
                    value=str(bai_chon.get("ghi_chu", "")) if pd.notna(bai_chon.get("ghi_chu", "")) else "",
                    height=70
                )

                btn_sua = st.form_submit_button("💾  Lưu sửa đổi", use_container_width=True)
                if btn_sua:
                    try:
                        supabase.table("dang_ky_tin_bai").update({
                            "tieu_de":   ten_moi.strip(),
                            "nguoi_gui": nguoi_gui_moi.strip(),
                            "duong_dan": link_moi.strip(),
                            "ghi_chu":   ghi_chu_moi.strip(),
                        }).eq("id", int(id_chon)).execute()
                        st.success("✅ Đã cập nhật thành công!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Lỗi: {e}")

    # ══ Khu vực nguy hiểm ══
    with st.expander("⚠️ Khu vực nguy hiểm – Reset toàn bộ dữ liệu", expanded=False):
        st.error("Thao tác dưới đây sẽ **xóa toàn bộ dữ liệu** và **không thể hoàn tác!**")
        xac_nhan = st.text_input("Gõ **XOA** để xác nhận:", placeholder="Nhập XOA để xác nhận…", key="confirm_reset")
        if st.button("🗑️ Xóa sạch toàn bộ tin bài (Reset)", key="btn_reset") and xac_nhan == "XOA":
            try:
                supabase.table("dang_ky_tin_bai").delete().neq("id", 0).execute()
                st.success("Đã dọn sạch toàn bộ dữ liệu!")
                st.rerun()
            except Exception as e:
                st.error(f"Lỗi: {e}")

# ══════════════════════════════════════════════
# TAB 3 – THỐNG KÊ & BIỂU ĐỒ
# ══════════════════════════════════════════════
with tab3:
    st.markdown('<p class="section-title">📊 Báo cáo thống kê tin bài</p>', unsafe_allow_html=True)
    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    res_all = supabase.table("dang_ky_tin_bai").select("*").execute()
    if not res_all.data:
        st.info("Hệ thống chưa có dữ liệu tin bài.")
    else:
        df_all = pd.DataFrame(res_all.data)
        df_all["ngay_dang_ky"] = pd.to_datetime(df_all["ngay_dang_ky"], errors="coerce")
        df_all = df_all.dropna(subset=["ngay_dang_ky"])
        df_all["Năm"] = df_all["ngay_dang_ky"].dt.year

        col_nam, col_kieu, col_ct = st.columns(3)
        with col_nam:
            ds_nam   = sorted(df_all["Năm"].unique().tolist(), reverse=True)
            chon_nam = st.selectbox("📅 Năm:", ds_nam)

        df_nam = df_all[df_all["Năm"] == chon_nam].copy()

        with col_kieu:
            kieu_loc = st.selectbox("🔍 Lọc theo:", ["Cả năm", "Theo Quý", "Theo Tháng", "Theo Tuần", "Theo Ngày"])

        df_loc = df_nam.copy()
        with col_ct:
            if kieu_loc == "Theo Quý":
                df_nam["Quý"] = df_nam["ngay_dang_ky"].dt.quarter
                q_opts = sorted(df_nam["Quý"].unique())
                if q_opts:
                    chon = st.selectbox("Chọn:", [f"Quý {q}" for q in q_opts])
                    df_loc = df_nam[df_nam["Quý"] == int(chon.split()[-1])]
            elif kieu_loc == "Theo Tháng":
                df_nam["Tháng"] = df_nam["ngay_dang_ky"].dt.month
                t_opts = sorted(df_nam["Tháng"].unique())
                if t_opts:
                    chon = st.selectbox("Chọn:", [f"Tháng {t}" for t in t_opts])
                    df_loc = df_nam[df_nam["Tháng"] == int(chon.split()[-1])]
            elif kieu_loc == "Theo Tuần":
                df_nam["Tuần"] = df_nam["ngay_dang_ky"].dt.isocalendar().week
                tu_opts = sorted(df_nam["Tuần"].unique())
                if tu_opts:
                    chon = st.selectbox("Chọn:", [f"Tuần thứ {t}" for t in tu_opts])
                    df_loc = df_nam[df_nam["Tuần"] == int(chon.split()[-1])]
            elif kieu_loc == "Theo Ngày":
                ng_opts = sorted(df_nam["ngay_dang_ky"].dt.date.unique(), reverse=True)
                if ng_opts:
                    chon = st.selectbox("Chọn:", ng_opts, format_func=lambda x: x.strftime("%d/%m/%Y"))
                    df_loc = df_nam[df_nam["ngay_dang_ky"].dt.date == chon]
            else:
                st.info(f"📌 Hiển thị toàn bộ năm {chon_nam}")

        st.markdown("---")
        if df_loc.empty:
            st.warning("Không có bài viết nào trong khoảng thời gian đã chọn.")
        else:
            so_tu_viet = len(df_loc[df_loc["nguon_tin"].str.contains("Viết mới|tự viết",  case=False, na=False)])
            so_suu_tam = len(df_loc[df_loc["nguon_tin"].str.contains("Sưu tầm|đăng lại", case=False, na=False)])
            so_fb      = int(df_loc.get("dang_facebook", pd.Series(dtype=bool)).fillna(False).sum()) if "dang_facebook" in df_loc else 0
            so_zalo    = int(df_loc.get("dang_zalo",    pd.Series(dtype=bool)).fillna(False).sum()) if "dang_zalo" in df_loc else 0

            m1, m2, m3, m4, m5 = st.columns(5)
            m1.metric("📰 Tổng bài",    len(df_loc))
            m2.metric("✏️ Tự viết",    so_tu_viet)
            m3.metric("🔗 Sưu tầm",    so_suu_tam)
            m4.metric("📘 Đăng FB",    so_fb)
            m5.metric("💬 Đăng Zalo",  so_zalo)

            st.markdown("")

            chart_col1, chart_col2 = st.columns([3, 2])

            with chart_col1:
                df_bd = df_loc.groupby(["nguoi_gui", "nguon_tin"]).size().reset_index(name="Số lượng")
                fig_bar = px.bar(
                    df_bd,
                    x="nguoi_gui", y="Số lượng",
                    color="nguon_tin",
                    barmode="group",
                    text="Số lượng",
                    color_discrete_sequence=["#0A2744", "#B22222"],
                    labels={"nguoi_gui": "Người gửi", "nguon_tin": "Loại bài"},
                )
                fig_bar.update_traces(textposition="outside", marker_line_width=0)
                fig_bar.update_layout(
                    title=dict(text="📊 Phân bổ tin bài theo Cán bộ", font=dict(size=14, color="#0A2744")),
                    plot_bgcolor="white", paper_bgcolor="white",
                    yaxis=dict(gridcolor="#EEF2F7"),
                    font=dict(family="Be Vietnam Pro, Segoe UI"),
                    legend_title="Loại tin bài",
                    margin=dict(t=50, b=40),
                )
                st.plotly_chart(fig_bar, use_container_width=True)

            with chart_col2:
                pie_data = pd.DataFrame({
                    "Loại": ["Bài tự viết", "Bài sưu tầm"],
                    "Số lượng": [so_tu_viet, so_suu_tam]
                })
                fig_pie = px.pie(
                    pie_data, names="Loại", values="Số lượng",
                    color_discrete_sequence=["#0A2744", "#B22222"],
                    hole=0.4,
                )
                fig_pie.update_traces(
                    textposition="inside", textinfo="percent+label",
                    textfont=dict(size=12, family="Be Vietnam Pro"),
                )
                fig_pie.update_layout(
                    title=dict(text="🥧 Tỷ lệ loại bài", font=dict(size=14, color="#0A2744")),
                    plot_bgcolor="white", paper_bgcolor="white",
                    font=dict(family="Be Vietnam Pro, Segoe UI"),
                    showlegend=True,
                    margin=dict(t=50, b=40),
                )
                st.plotly_chart(fig_pie, use_container_width=True)

            # ── Xu hướng theo thời gian ──
            if kieu_loc in ("Cả năm", "Theo Quý"):
                df_loc["Tháng"] = df_loc["ngay_dang_ky"].dt.to_period("M").astype(str)
                df_trend = df_loc.groupby("Tháng").size().reset_index(name="Số bài")
                if len(df_trend) > 1:
                    fig_line = px.line(
                        df_trend, x="Tháng", y="Số bài",
                        markers=True,
                        color_discrete_sequence=["#0A2744"],
                        labels={"Tháng": "Tháng", "Số bài": "Số lượng bài"},
                    )
                    fig_line.update_traces(
                        line=dict(width=2.5),
                        marker=dict(size=8, color="#B22222", symbol="circle"),
                    )
                    fig_line.update_layout(
                        title=dict(text="📈 Xu hướng tin bài theo tháng", font=dict(size=14, color="#0A2744")),
                        plot_bgcolor="white", paper_bgcolor="white",
                        yaxis=dict(gridcolor="#EEF2F7"),
                        font=dict(family="Be Vietnam Pro, Segoe UI"),
                        margin=dict(t=50, b=40),
                    )
                    st.plotly_chart(fig_line, use_container_width=True)
